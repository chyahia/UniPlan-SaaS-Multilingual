# Copyright (c) 2026 Chaib Yahia. All rights reserved.
# This software is licensed under the CC BY-NC 4.0 License. Commercial use is strictly prohibited.
from flask import Blueprint, request, jsonify, send_file, session
import json
import io
from datetime import datetime
from collections import defaultdict
from app.database import db, ExamSetting, ExamTeacher, ExamSubject, ExamLevel, ExamRoom

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import pandas as pd
from app.services.exams_algorithms import _run_initial_subject_placement, clean_string_for_matching
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment
from flask_babel import _

exams_export_bp = Blueprint('exams_export', __name__)

# ================== قاموس الترجمة للامتحانات ==================
EXAM_TRANSLATIONS = {
    'ar': {
        'period': _('الفترة'), 'date_day': _('اليوم/التاريخ'), 'exam_schedule': _('جدول امتحانات'),
        'guard_schedule': _('جدول الحراسة'), 'simplified': _('(مُبسَّط)'),
        'course_prof': _('أستاذ المادة'), 'guarding': _('الحراسة:'), 'large_hall': _('القاعة الكبيرة'),
        'other_halls': _('القاعات الأخرى'), 'not_specified': _('غير محدد'), 'no_level': _('بدون مستوى'),
        'no_guards': _('(لا يوجد)'), 'shortage_alert': _('نقص!'), 'empty': _('- فراغ -'),
        'is_guarding': _('(حراسة)'), 'no_guarding': _('(دون حراسة)'), 'assigned_guard': _('(تكليف بحراسة)'),
        'days': [_("الأحد"), _("الاثنين"), _("الثلاثاء"), _("الأربعاء"), _("الخميس"), _("الجمعة"), _("السبت")]
    },
    'en': {
        'period': 'Time Slot', 'date_day': 'Day / Date', 'exam_schedule': 'Exams Schedule',
        'guard_schedule': 'Guarding Schedule', 'simplified': '(Simplified)',
        'course_prof': 'Course Professor', 'guarding': 'Invigilators:', 'large_hall': 'Large Hall',
        'other_halls': 'Other Halls', 'not_specified': 'Not Specified', 'no_level': 'No Level',
        'no_guards': '(None)', 'shortage_alert': 'Shortage!', 'empty': '- Empty -',
        'is_guarding': '(Invigilation)', 'no_guarding': '(No Invigilation)', 'assigned_guard': '(Assigned Invigilator)',
        'days': ["Sunday", "Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday"]
    },
    'fr': {
        'period': 'Créneau', 'date_day': 'Jour / Date', 'exam_schedule': 'Emploi des Examens',
        'guard_schedule': 'Emploi de Surveillance', 'simplified': '(Simplifié)',
        'course_prof': 'Professeur du module', 'guarding': 'Surveillants:', 'large_hall': 'Grande Salle',
        'other_halls': 'Autres Salles', 'not_specified': 'Non Spécifié', 'no_level': 'Sans Niveau',
        'no_guards': '(Aucun)', 'shortage_alert': 'Manque!', 'empty': '- Vide -',
        'is_guarding': '(Surveillance)', 'no_guarding': '(Sans Surveillance)', 'assigned_guard': '(Assigné à surveiller)',
        'days': ["Dimanche", "Lundi", "Mardi", "Mercredi", "Jeudi", "Vendredi", "Samedi"]
    }
}

def create_word_document_with_table(doc, title, headers, data_grid, lang='ar'):
    heading = doc.add_heading(level=2)
    heading.clear() 
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if lang == 'ar':
        pPr = heading._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
    
    run = heading.add_run(title)
    font = run.font
    if lang == 'ar': font.rtl = True
    font.name = 'Arial'

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = False

    if lang == 'ar':
        tbl_pr = table._element.xpath('w:tblPr')[0]
        bidi_visual_element = OxmlElement('w:bidiVisual')
        tbl_pr.append(bidi_visual_element)

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell_paragraph = hdr_cells[i].paragraphs[0]
        cell_paragraph.text = ""
        run = cell_paragraph.add_run(header)
        font = run.font
        if lang == 'ar': font.rtl = True
        font.name = 'Arial'
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if lang == 'ar': cell_paragraph.paragraph_format.rtl = True

    for row_data in data_grid:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            cell_paragraph = row_cells[i].paragraphs[0]
            cell_paragraph.text = ""
            lines = str(cell_data).split('\n')
            for idx, line in enumerate(lines):
                if idx > 0:
                    cell_paragraph.add_run().add_break()
                run = cell_paragraph.add_run(line)
                font = run.font
                if lang == 'ar': font.rtl = True
                font.name = 'Arial'
            cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if lang == 'ar' else WD_ALIGN_PARAGRAPH.LEFT
            if lang == 'ar': cell_paragraph.paragraph_format.rtl = True
            
    doc.add_page_break()

# ==============================================================
# 1. تصدير جداول الامتحانات لكل المستويات
# ==============================================================
@exams_export_bp.route('/exams/api/export/word/all-exams', methods=['POST'])
def export_exams_word():
    tenant_id = session.get('tenant_id')
    if not tenant_id: return jsonify({"error": _("غير مصرح")}), 403

    schedule_data = request.get_json()
    if not schedule_data: return jsonify({"error": "No schedule data provided"}), 400
    
    lang = request.args.get('lang', 'ar')
    t = EXAM_TRANSLATIONS.get(lang, EXAM_TRANSLATIONS['ar'])
    
    assignments_rows = []
    for t_obj in ExamTeacher.query.filter_by(tenant_id=tenant_id).all():
        for s in t_obj.subjects:
            levels_list = sorted([l.name for l in s.levels]) if hasattr(s, 'levels') and s.levels else []
            combined_level = " + ".join(levels_list) if levels_list else t['no_level']
            assignments_rows.append({
                'subj_name': s.name, 
                'level_name': combined_level, 
                'prof_name': t_obj.name
            })
            
    settings_row = ExamSetting.query.filter_by(key='main_settings', tenant_id=tenant_id).first()
    settings_data = json.loads(settings_row.value) if settings_row and settings_row.value else {}
    
    guards_large = int(settings_data.get('guardsLargeHall', 4))
    guards_medium = int(settings_data.get('guardsMediumHall', 2))
    guards_small = int(settings_data.get('guardsSmallHall', 1))

    subject_owners = {(row['subj_name'], row['level_name']): row['prof_name'] for row in assignments_rows}
    
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin = Cm(0.5)
    section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = margin, margin, margin, margin

    all_dates = sorted(schedule_data.keys())
    all_times = sorted({time for date_slots in schedule_data.values() for time in date_slots})
    all_levels = sorted({exam['level'] for slots in schedule_data.values() for exams in slots.values() for exam in exams})
    day_names = t['days']
    
    headers = [t['period']] + [f"{day_names[datetime.strptime(d, '%Y-%m-%d').isoweekday() % 7]}\n{d}" for d in all_dates]

    for level in all_levels:
        data_grid = []
        for time in all_times:
            row_data = [time]
            for date in all_dates:
                exam = next((e for e in schedule_data.get(date, {}).get(time, []) if e['level'] == level), None)
                content = ""
                if exam:
                    owner = subject_owners.get((exam['subject'], exam['level']), t['not_specified'])
                    content = f"{exam['subject']}\n{t['course_prof']}: {owner}\n\n{t['guarding']}"

                    halls_by_type = defaultdict(list)
                    for h in exam.get('halls', []): halls_by_type[h['type']].append(h['name'])
                    
                    # ✨ التعديل 1: تصفية كلمة "نقص" بذكاء لتشمل الصيغة العربية والإنجليزية معاً
                    guards_copy = [g for g in exam.get('guards', []) if g not in ["**نقص**", _("**نقص**")]]

                    # ✨ التعديل 2: استخراج القاعات دون الاعتماد على لغة الجلسة فقط لتفادي مشكلة الترجمة
                    large_halls = []
                    other_hall_names = []
                    
                    for h_type, h_names in halls_by_type.items():
                        if h_type in ['كبيرة', _('كبيرة'), 'Large', 'large']:
                            large_halls.extend(h_names)
                        else:
                            other_hall_names.extend(h_names)

                    if large_halls:
                        num_guards_needed = len(large_halls) * guards_large
                        g_list = guards_copy[:num_guards_needed]
                        guards_copy = guards_copy[num_guards_needed:]
                        hall_names = ", ".join(large_halls)
                        guard_text = '\n'.join(g_list) if g_list else t['no_guards']
                        content += f"\n{t['large_hall']}: {hall_names}\n{guard_text}"
                    
                    if other_hall_names:
                        guard_text = '\n'.join(guards_copy) if guards_copy else t['no_guards']
                        content += f"\n{t['other_halls']}: {', '.join(other_hall_names)}\n{guard_text}"
                
                row_data.append(content)
            data_grid.append(row_data)
        
        create_word_document_with_table(doc, f"{t['exam_schedule']}: {level}", headers, data_grid, lang)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="export.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ==============================================================
# 2. تصدير جداول الحراسة المفصلة للأساتذة
# ==============================================================
@exams_export_bp.route('/exams/api/export/word/all-profs', methods=['POST'])
def export_profs_word():
    tenant_id = session.get('tenant_id')
    if not tenant_id: return jsonify({"error": _("غير مصرح")}), 403

    schedule_data = request.get_json()
    if not schedule_data: return jsonify({"error": "No schedule data provided"}), 400

    lang = request.args.get('lang', 'ar')
    t = EXAM_TRANSLATIONS.get(lang, EXAM_TRANSLATIONS['ar'])

    all_professors = sorted([p.name for p in ExamTeacher.query.filter_by(tenant_id=tenant_id).all()])
    
    prof_owned_subjects = defaultdict(set)
    for t_obj in ExamTeacher.query.filter_by(tenant_id=tenant_id).all():
        for s in t_obj.subjects:
            levels_list = sorted([l.name for l in s.levels]) if hasattr(s, 'levels') and s.levels else []
            combined_level = " + ".join(levels_list) if levels_list else t['no_level']
            prof_owned_subjects[t_obj.name].add((s.name, combined_level))

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin = Cm(0.5)
    section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = margin, margin, margin, margin

    all_dates = sorted(schedule_data.keys())
    all_times = sorted({time for date_slots in schedule_data.values() for time in date_slots})
    day_names = t['days']
    
    for prof_name in all_professors:
        title = f"{t['guard_schedule']}: {prof_name}"
        headers = [t['date_day']] + all_times
        
        heading = doc.add_heading(level=2); heading.clear()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if lang == 'ar':
            pPr = heading._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)
        run = heading.add_run(title)
        font = run.font; 
        if lang == 'ar': font.rtl = True
        font.name = 'Arial'

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = False
        
        if lang == 'ar':
            tbl_pr = table._element.xpath('w:tblPr')[0]
            bidi_visual_element = OxmlElement('w:bidiVisual')
            tbl_pr.append(bidi_visual_element)
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]; p.text = ""
            run = p.add_run(header)
            font = run.font; 
            if lang == 'ar': font.rtl = True
            font.name = 'Arial'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if lang == 'ar': p.paragraph_format.rtl = True

        has_any_duty = False
        for date in all_dates:
            row_cells = table.add_row().cells
            day_name = day_names[datetime.strptime(date, '%Y-%m-%d').isoweekday() % 7]
            
            p = row_cells[0].paragraphs[0]; p.text = ""
            run = p.add_run(f"{day_name}\n{date}"); 
            if lang == 'ar': run.font.rtl = True
            run.font.name = 'Arial'; run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; 
            if lang == 'ar': p.paragraph_format.rtl = True

            for i, time in enumerate(all_times, 1):
                cell_content_parts = []
                is_teaching_and_guarding = False
                is_teaching_only = False
                
                exams_in_slot = schedule_data.get(date, {}).get(time, [])
                
                for exam in exams_in_slot:
                    is_guarding = prof_name in exam.get('guards', [])
                    is_owner = (exam['subject'], exam['level']) in prof_owned_subjects.get(prof_name, set())

                    if is_guarding or is_owner:
                        has_any_duty = True
                        if is_guarding:
                            if is_owner: is_teaching_and_guarding = True
                            cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n{t['is_guarding']}")
                        elif is_owner:
                            is_teaching_only = True
                            cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n{t['no_guarding']}")
                
                p = row_cells[i].paragraphs[0]; p.text = ""
                lines = "\n---\n".join(cell_content_parts).split('\n')
                for idx, line in enumerate(lines):
                    if idx > 0: p.add_run().add_break()
                    run = p.add_run(line)
                    font = run.font; 
                    if lang == 'ar': font.rtl = True
                    font.name = 'Arial'
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if lang == 'ar' else WD_ALIGN_PARAGRAPH.LEFT
                if lang == 'ar': p.paragraph_format.rtl = True
                
                shading_elm = OxmlElement('w:shd')
                if is_teaching_and_guarding:
                    shading_elm.set(qn('w:fill'), 'D4EDDA') # تظليل أخضر خفيف
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
                elif is_teaching_only:
                    shading_elm.set(qn('w:fill'), 'FFF3CD') # تظليل أصفر خفيف
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        if has_any_duty:
             doc.add_page_break()
        else:
            doc._body.remove(table._element)
            doc._body.remove(heading._element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="export.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ==============================================================
# 3. تصدير جداول الحراسة المبسطة للأساتذة
# ==============================================================
@exams_export_bp.route('/exams/api/export/word/all-profs-anonymous', methods=['POST'])
def export_profs_anonymous_word():
    tenant_id = session.get('tenant_id')
    if not tenant_id: return jsonify({"error": _("غير مصرح")}), 403

    schedule_data = request.get_json()
    if not schedule_data: return jsonify({"error": "No schedule data provided"}), 400

    lang = request.args.get('lang', 'ar')
    t = EXAM_TRANSLATIONS.get(lang, EXAM_TRANSLATIONS['ar'])

    all_professors = sorted([p.name for p in ExamTeacher.query.filter_by(tenant_id=tenant_id).all()])
    
    prof_owned_subjects = defaultdict(set)
    for t_obj in ExamTeacher.query.filter_by(tenant_id=tenant_id).all():
        for s in t_obj.subjects:
            levels_list = sorted([l.name for l in s.levels]) if hasattr(s, 'levels') and s.levels else []
            combined_level = " + ".join(levels_list) if levels_list else t['no_level']
            prof_owned_subjects[t_obj.name].add((s.name, combined_level))

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin = Cm(0.5)
    section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = margin, margin, margin, margin

    all_dates = sorted(schedule_data.keys())
    all_times = sorted({time for date_slots in schedule_data.values() for time in date_slots})
    day_names = t['days']
    
    for prof_name in all_professors:
        title = f"{t['guard_schedule']} {t['simplified']}: {prof_name}"
        headers = [t['date_day']] + all_times
        
        heading = doc.add_heading(level=2); heading.clear()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if lang == 'ar':
            pPr = heading._p.get_or_add_pPr()
            bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)
        run = heading.add_run(title)
        font = run.font; 
        if lang == 'ar': font.rtl = True
        font.name = 'Arial'

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = False
        
        if lang == 'ar':
            tbl_pr = table._element.xpath('w:tblPr')[0]
            bidi_visual_element = OxmlElement('w:bidiVisual')
            tbl_pr.append(bidi_visual_element)
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]; p.text = ""
            run = p.add_run(header)
            font = run.font; 
            if lang == 'ar': font.rtl = True
            font.name = 'Arial'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if lang == 'ar': p.paragraph_format.rtl = True

        has_any_duty = False
        for date in all_dates:
            row_cells = table.add_row().cells
            day_name = day_names[datetime.strptime(date, '%Y-%m-%d').isoweekday() % 7]
            
            p = row_cells[0].paragraphs[0]; p.text = ""
            run = p.add_run(f"{day_name}\n{date}"); 
            if lang == 'ar': run.font.rtl = True
            run.font.name = 'Arial'; run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; 
            if lang == 'ar': p.paragraph_format.rtl = True

            for i, time in enumerate(all_times, 1):
                cell_content_parts = []
                is_teaching_and_guarding = False
                is_teaching_only = False
                
                exams_in_slot = schedule_data.get(date, {}).get(time, [])
                
                for exam in exams_in_slot:
                    is_guarding = prof_name in exam.get('guards', [])
                    is_owner = (exam['subject'], exam['level']) in prof_owned_subjects.get(prof_name, set())

                    if is_guarding or is_owner:
                        has_any_duty = True
                        if is_guarding:
                            if is_owner:
                                is_teaching_and_guarding = True
                                cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n{t['is_guarding']}")
                            else:
                                cell_content_parts.append(t['assigned_guard'])
                        elif is_owner:
                            is_teaching_only = True
                            cell_content_parts.append(f"{exam['subject']} ({exam['level']})\n{t['no_guarding']}")
                
                p = row_cells[i].paragraphs[0]; p.text = ""
                lines = "\n---\n".join(cell_content_parts).split('\n')
                for idx, line in enumerate(lines):
                    if idx > 0: p.add_run().add_break()
                    run = p.add_run(line)
                    font = run.font; 
                    if lang == 'ar': font.rtl = True
                    font.name = 'Arial'
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if lang == 'ar' else WD_ALIGN_PARAGRAPH.LEFT
                if lang == 'ar': p.paragraph_format.rtl = True
                
                shading_elm = OxmlElement('w:shd')
                if is_teaching_and_guarding:
                    shading_elm.set(qn('w:fill'), 'D4EDDA')
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
                elif is_teaching_only:
                    shading_elm.set(qn('w:fill'), 'FFF3CD')
                    row_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        if has_any_duty:
             doc.add_page_break()
        else:
            doc._body.remove(table._element)
            doc._body.remove(heading._element)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="export.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ==============================================================
# 4. تصدير مخطط التوزيع اليدوي (قالب إكسل)
# ==============================================================
@exams_export_bp.route('/exams/api/export-manual-distribution-template', methods=['POST'])
def export_manual_distribution_template():
    tenant_id = session.get('tenant_id')
    if not tenant_id: return jsonify({"error": _("غير مصرح")}), 403

    try:
        # ✨ 1. استخراج لغة الجلسة الحالية لتحديد اتجاه الملف
        lang = session.get('lang', 'ar')
        is_rtl = (lang == 'ar')

        row_sched = ExamSetting.query.filter_by(key='exam_schedule', tenant_id=tenant_id).first()
        exam_schedule = json.loads(row_sched.value) if row_sched and row_sched.value else {}
        
        all_dates = sorted(exam_schedule.keys())
        all_times = sorted(list(set(time for slots in exam_schedule.values() for slot in slots for time in [slot['time']])))
        
        if not all_dates or not all_times:
            return jsonify({"error": _("الرجاء حفظ جدول الأيام والفترات في المرحلة 4 أولاً لتتمكن من تصدير المخطط.")}), 400

        all_levels_list = [l.name for l in ExamLevel.query.filter_by(tenant_id=tenant_id).all()]
        
        all_subjects = []
        original_subject_map = {} 
        for s in ExamSubject.query.filter_by(tenant_id=tenant_id).all():
            levels_list = sorted([l.name for l in s.levels]) if hasattr(s, 'levels') and s.levels else []
            combined_level = " + ".join(levels_list) if levels_list else _("بدون مستوى")
            all_subjects.append({'name': s.name, 'level': combined_level, 'levels': levels_list})
            
            c_name = clean_string_for_matching(s.name)
            c_levels = tuple(sorted([clean_string_for_matching(l) for l in levels_list]))
            original_subject_map[(c_name, c_levels)] = s.name
            
        all_halls = [{'name': h.name, 'type': h.type} for h in ExamRoom.query.filter_by(tenant_id=tenant_id).all()]
        
        subject_owners = {}
        for t in ExamTeacher.query.filter_by(tenant_id=tenant_id).all():
            for s in t.subjects:
                levels_list = sorted([l.name for l in s.levels]) if hasattr(s, 'levels') and s.levels else []
                if levels_list:
                    levels_tuple = tuple(sorted([clean_string_for_matching(l) for l in levels_list]))
                    subject_owners[(clean_string_for_matching(s.name), levels_tuple)] = t.name

        level_hall_assignments = defaultdict(list)
        for l in ExamLevel.query.filter_by(tenant_id=tenant_id).all():
            for r in l.rooms:
                level_hall_assignments[l.name].append(r.name)

        settings_for_placement = {
            'examSchedule': exam_schedule,
            'levelHallAssignments': dict(level_hall_assignments)
        }

        # تم تصحيح خطأ طمس المتغير هنا مسبقاً (استخدام _ignored_stats)
        initial_schedule, _ignored_stats = _run_initial_subject_placement(settings_for_placement, all_subjects, all_levels_list, subject_owners, all_halls)

        output = io.BytesIO()
        writer = pd.ExcelWriter(output, engine='openpyxl')
        import re 
        
        if not all_levels_list:
            df_empty = pd.DataFrame([_("لا توجد مستويات مدخلة بعد")])
            df_empty.to_excel(writer, sheet_name=_("فارغ"))
            
        for level_name in sorted(all_levels_list):
            df_level = pd.DataFrame(index=all_times, columns=all_dates)
            
            # ✨ 2. ترجمة ترويسة الفترة ديناميكياً
            df_level.index.name = "Time Slot" if lang == 'en' else _("الفترة")
            c_level_name = clean_string_for_matching(level_name)
            
            for date, slots in initial_schedule.items():
                for time, exams in slots.items():
                    for exam in exams:
                        if c_level_name in exam.get('levels_list', []):
                            exam_c_name = exam['subject']
                            exam_c_levels = tuple(sorted(exam.get('levels_list', [])))
                            orig_subj_name = original_subject_map.get((exam_c_name, exam_c_levels), exam_c_name)
                            
                            cell_content = f"{orig_subj_name}\n::: {exam['professor']}\n::: {exam['level']}"
                            
                            existing = df_level.at[time, date]
                            if pd.notna(existing) and str(existing).strip() != '':
                                df_level.at[time, date] = str(existing) + "\n\n====================\n\n" + cell_content
                            else:
                                df_level.at[time, date] = cell_content
            
            unplaced_subjects = []
            for s in all_subjects:
                if level_name in s.get('levels', []):
                    is_placed = False
                    c_s_name = clean_string_for_matching(s['name'])
                    c_s_level = clean_string_for_matching(s['level'])
                    for d in initial_schedule.values():
                        for t in d.values():
                            for e in t:
                                if e['subject'] == c_s_name and clean_string_for_matching(e['level']) == c_s_level:
                                    is_placed = True
                    if not is_placed:
                        unplaced_subjects.append(s)

            if unplaced_subjects:
                # ✨ ترجمة صف "المواد غير الموزعة" للإنجليزية
                unplaced_row_name = "--- Unassigned Subjects ---" if lang == 'en' else _("--- مواد غير موزعة ---")
                df_level.loc[unplaced_row_name] = ''
                cell_texts = []
                for s in unplaced_subjects:
                    s_tuple = tuple(sorted([clean_string_for_matching(l) for l in s.get('levels', [])]))
                    owner = subject_owners.get((clean_string_for_matching(s['name']), s_tuple), _('غير محدد'))
                    cell_texts.append(f"{s['name']}\n::: {owner}\n::: {s['level']}")
                if all_dates:
                    df_level.at[unplaced_row_name, all_dates[0]] = "\n\n====================\n\n".join(cell_texts)

            safe_sheet_name = re.sub(r'[\\*?:/\[\]]', '-', level_name)[:31]
            df_level.to_excel(writer, sheet_name=safe_sheet_name)
            worksheet = writer.sheets[safe_sheet_name]

            # ✨ 3. التحكم الديناميكي في اتجاه الورقة حسب اللغة (يمين-يسار للعربية، ويسار-يمين للإنجليزية)
            worksheet.sheet_view.rightToLeft = is_rtl
            worksheet.column_dimensions['A'].width = 18
            for i in range(2, len(all_dates) + 2):
                worksheet.column_dimensions[get_column_letter(i)].width = 25
                
            # ✨ 4. تحديد محاذاة النص واتجاه القراءة (Reading Order) بناءً على اللغة
            align_horizontal = 'right' if is_rtl else 'left'
            reading_order = 2 if is_rtl else 1 
            wrap_alignment = Alignment(wrap_text=True, horizontal=align_horizontal, vertical='center', readingOrder=reading_order)
            
            for row in worksheet.iter_rows():
                if row[0].row == 1:
                    worksheet.row_dimensions[row[0].row].height = 35 
                else:
                    worksheet.row_dimensions[row[0].row].height = None
                
                for cell in row:
                    cell.alignment = wrap_alignment

        writer.close()
        
        # ✨ 5. ترجمة اسم الملف المصدر
        file_name = "Manual_Distribution_Template.xlsx" if lang == 'en' else _('مخطط_توزيع_المواد_للتعديل.xlsx')
        
        return send_file(
            io.BytesIO(output.getvalue()), 
            as_attachment=True, 
            download_name=file_name, 
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# ==============================================================
# 5. استيراد المخطط اليدوي المُعدل
# ==============================================================
@exams_export_bp.route('/exams/api/import-manual-distribution', methods=['POST'])
def import_manual_distribution():
    tenant_id = session.get('tenant_id')
    if not tenant_id: return jsonify({"error": _("غير مصرح")}), 403

    if 'file' not in request.files: return jsonify({"error": _("لم يتم العثور على ملف.")}), 400
    file = request.files['file']
    try:
        xls = pd.read_excel(file, sheet_name=None, index_col=0, dtype=str)
        pinned_schedule = defaultdict(lambda: defaultdict(list))
        level_hall_assignments = defaultdict(list)
        for l in ExamLevel.query.filter_by(tenant_id=tenant_id).all():
            for r in l.rooms: level_hall_assignments[l.name].append({'name': r.name, 'type': r.type})

        pinned_count = 0
        for sheet_name, df in xls.items():
            for date in df.columns:
                for time in df.index:
                    cell_value = df.at[time, date]
                    if pd.notna(cell_value):
                        cell_str = str(cell_value).strip()
                        # ذكاء اصطناعي لقراءة المربعات المنسقة أو القديمة
                        if "====================" in cell_str:
                            exams_in_cell = cell_str.split('\n\n====================\n\n')
                        elif "\n:::" in cell_str:
                            exams_in_cell = [cell_str]
                        else:
                            exams_in_cell = cell_str.split('\n')

                        for exam_block in exams_in_cell:
                            clean_block = exam_block.replace('\n', ' ')
                            if ':::' in clean_block:
                                try:
                                    parts = [part.strip() for part in clean_block.split(':::')]
                                    
                                    # ✨ التعديل هنا: تحويل الوقت لنص، وتجاهل صف "المواد غير الموزعة" باللغتين العربية والإنجليزية
                                    time_str = str(time)
                                    if len(parts) >= 3 and date and time_str and "مواد غير موزعة" not in time_str and "Unassigned Subjects" not in time_str:
                                        
                                        subject_name, professor_name, level_name = parts[0], parts[1], parts[2]
                                        halls_details = level_hall_assignments.get(level_name, [])
                                        levels_list = level_name.split(' + ') if ' + ' in level_name else [level_name]
                                        
                                        exam = {
                                            "date": str(date).strip(), "time": time_str.strip(),
                                            "subject": subject_name, "level": level_name,
                                            "levels_list": [clean_string_for_matching(l) for l in levels_list],
                                            "professor": professor_name, "halls": halls_details, "guards": []
                                        }
                                        pinned_schedule[exam['date']][exam['time']].append(exam)
                                        pinned_count += 1
                                except ValueError: continue
        
        setting = ExamSetting.query.filter_by(key='pinned_subject_schedule', tenant_id=tenant_id).first()
        value_str = json.dumps(pinned_schedule)
        if setting: setting.value = value_str
        else: db.session.add(ExamSetting(key='pinned_subject_schedule', value=value_str, tenant_id=tenant_id))
        db.session.commit()
        return jsonify({"success": True, "message": _("تم استيراد وتثبيت {pinned_count} مادة بنجاح.").format(pinned_count=pinned_count)})
    except Exception as e: 
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# ==============================================================
# 6. مسح المخطط اليدوي
# ==============================================================
@exams_export_bp.route('/exams/api/clear-manual-distribution', methods=['POST'])
def clear_manual_distribution():
    tenant_id = session.get('tenant_id')
    if not tenant_id: return jsonify({"error": _("غير مصرح")}), 403

    try:
        setting = ExamSetting.query.filter_by(key='pinned_subject_schedule', tenant_id=tenant_id).first()
        if setting:
            db.session.delete(setting)
            db.session.commit()
        return jsonify({"success": True, "message": _("تم مسح الجدول اليدوي. سيعتمد التشغيل القادم على التوزيع التلقائي.")})
    except Exception as e: db.session.rollback(); return jsonify({"error": str(e)}), 500

# ==============================================================
# 7. تصدير الجدول النهائي (مع الحراس) للإكسل
# ==============================================================
@exams_export_bp.route('/exams/api/export-final-excel', methods=['POST'])
def export_final_excel():
    tenant_id = session.get('tenant_id')
    if not tenant_id: return jsonify({"error": _("غير مصرح")}), 403

    schedule_data = request.json
    if not schedule_data: return jsonify({"error": _("لا توجد بيانات للجدول. قم بالتوليد أولاً.")}), 400

    try:
        import pandas as pd
        import io
        import re
        from openpyxl.utils import get_column_letter
        from openpyxl.styles import Alignment, Font

        # ✨ تحديد لغة الجلسة الحالية
        lang = session.get('lang', 'ar')
        is_rtl = (lang == 'ar')

        output = io.BytesIO()
        writer = pd.ExcelWriter(output, engine='openpyxl')
        
        all_dates = sorted(schedule_data.keys())
        all_times = sorted({time for date_slots in schedule_data.values() for time in date_slots})
        all_levels = sorted({exam['level'] for slots in schedule_data.values() for exams in slots.values() for exam in exams})
        
        # ✨ 1. قاموس ترجمة الأيام لترويسة الأعمدة
        day_trans = {
            "الأحد": "Sunday", "الإثنين": "Monday", "الاثنين": "Monday",
            "الثلاثاء": "Tuesday", "الأربعاء": "Wednesday",
            "الخميس": "Thursday", "الجمعة": "Friday", "السبت": "Saturday"
        }

        display_dates = []
        for d in all_dates:
            if lang == 'en':
                new_d = d
                for ar, en in day_trans.items():
                    if ar in new_d:
                        new_d = new_d.replace(ar, f"{ar} / {en}")
                display_dates.append(new_d)
            else:
                display_dates.append(d)

        # ✨ 2. تجهيز المفاتيح لتوضيح البيانات داخل الخانة
        subject_lbl = "مادة / Subject:" if lang == 'en' else "مادة:"
        prof_lbl = "أستاذ / Professor:" if lang == 'en' else "أستاذ:"
        level_lbl = "مستوى / Level:" if lang == 'en' else "مستوى:"
        halls_lbl = "قاعات / Halls:" if lang == 'en' else "قاعات:"
        guards_lbl = "حراس / Guards:" if lang == 'en' else "حراس:"

        for level_name in all_levels:
            df_level = pd.DataFrame(index=all_times, columns=display_dates)
            df_level.index.name = "الوقت / Time" if lang == 'en' else _("الفترة")
            
            for date_idx, date in enumerate(all_dates):
                disp_date = display_dates[date_idx]
                if date in schedule_data:
                    for time, exams in schedule_data[date].items():
                        for exam in exams:
                            if exam['level'] == level_name:
                                halls_str = "، ".join([h['name'] for h in exam.get('halls', [])])
                                guards_str = "، ".join([g for g in exam.get('guards', [])])
                                
                                # ✨ التعديل هنا: إضافة نزول سطر (\n) ومسافة بعد كل مفتاح لتظهر بشكل مرتب
                                cell_content = (
                                    f"{subject_lbl}\n {exam['subject']}\n"
                                    f"::: {prof_lbl}\n {exam['professor']}\n"
                                    f"::: {level_lbl}\n {exam['level']}\n"
                                    f"::: {halls_lbl}\n {halls_str}\n"
                                    f"::: {guards_lbl}\n {guards_str}"
                                )
                                
                                existing = df_level.at[time, disp_date]
                                if pd.notna(existing) and str(existing).strip() != '':
                                    df_level.at[time, disp_date] = str(existing) + "\n\n====================\n\n" + cell_content
                                else:
                                    df_level.at[time, disp_date] = cell_content
            
            safe_sheet_name = re.sub(r'[\\*?:/\[\]]', '-', level_name)[:31]
            df_level.to_excel(writer, sheet_name=safe_sheet_name)
            
            worksheet = writer.sheets[safe_sheet_name]
            worksheet.sheet_view.rightToLeft = is_rtl
            worksheet.column_dimensions['A'].width = 18
            for i in range(2, len(all_dates) + 2):
                worksheet.column_dimensions[get_column_letter(i)].width = 35
                
            # ✨ 3. محاذاة النص بناءً على اللغة
            align_horizontal = 'right' if is_rtl else 'left'
            reading_order = 2 if is_rtl else 1
            wrap_alignment = Alignment(wrap_text=True, horizontal=align_horizontal, vertical='center', readingOrder=reading_order)
            
            for row in worksheet.iter_rows():
                if row[0].row == 1:
                    worksheet.row_dimensions[row[0].row].height = 35
                else:
                    worksheet.row_dimensions[row[0].row].height = None
                
                for cell in row:
                    cell.alignment = wrap_alignment

            # ✨ 4. إضافة تحذير بأسفل الجدول لحماية المفاتيح
            last_row = worksheet.max_row + 2
            max_col = len(all_dates) + 1
            worksheet.merge_cells(start_row=last_row, start_column=1, end_row=last_row, end_column=max_col)
            note_cell = worksheet.cell(row=last_row, column=1)
            
            if lang == 'en':
                note_cell.value = "⚠️ Note: Do not change the keys before the colon (:). Only modify data after it."
            else:
                note_cell.value = "⚠️ ملاحظة هامة: الرجاء عدم تغيير الكلمات المكتوبة قبل النقطتين الرأسيتين (:). قم بتعديل البيانات بعدها فقط."
                
            note_cell.font = Font(bold=True, color="C00000")
            note_cell.alignment = Alignment(horizontal='center', vertical='center')

        writer.close()
        
        # تغيير اسم الملف بناءً على اللغة
        file_name = "Final_Invigilation_Schedule.xlsx" if lang == 'en' else _('الجدول_النهائي_للحراسة.xlsx')
        return send_file(io.BytesIO(output.getvalue()), as_attachment=True, download_name=file_name, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# ==============================================================
# 8. استيراد الجدول النهائي المُعدل ونشره
# ==============================================================
@exams_export_bp.route('/exams/api/import-final-excel', methods=['POST'])
def import_final_excel():
    tenant_id = session.get('tenant_id')
    if not tenant_id: return jsonify({"error": _("غير مصرح")}), 403

    if 'file' not in request.files: return jsonify({"error": _("لم يتم العثور على ملف.")}), 400
    file = request.files['file']
    
    try:
        import pandas as pd
        import re
        from collections import defaultdict
        import json
        
        xls = pd.read_excel(file, sheet_name=None, index_col=0, dtype=str)
        final_schedule = defaultdict(lambda: defaultdict(list))
        
        # ✨ دالة مساعدة لاستخلاص البيانات الصافية متجاهلة المفاتيح
        def extract_val(text):
            if ':' in text:
                return text.split(':', 1)[1].strip()
            return text.strip()
        
        for sheet_name, df in xls.items():
            for excel_date in df.columns:
                # ✨ تنظيف الترويسة من أسماء الأيام الإنجليزية لإعادتها للصيغة الأصلية
                clean_date = str(excel_date)
                clean_date = re.sub(r'\s*/\s*[A-Za-z]+', '', clean_date).strip()

                for time in df.index:
                    cell_value = df.at[time, excel_date]
                    if pd.notna(cell_value):
                        cell_str = str(cell_value).strip()
                        
                        if "====================" in cell_str:
                            exams_in_cell = cell_str.split('\n\n====================\n\n')
                        elif "\n:::" in cell_str:
                            exams_in_cell = [cell_str]
                        else:
                            exams_in_cell = cell_str.split('\n')
                            
                        for exam_block in exams_in_cell:
                            clean_block = exam_block.replace('\n', ' ')
                            if ':::' in clean_block:
                                parts = [part.strip() for part in clean_block.split(':::')]
                                if len(parts) >= 5:
                                    
                                    # ✨ قراءة القيم عبر الدالة المساعدة (آمنة مع الملفات القديمة والجديدة)
                                    subject_val = extract_val(parts[0])
                                    prof_val = extract_val(parts[1])
                                    level_val = extract_val(parts[2])
                                    halls_val = extract_val(parts[3])
                                    guards_val = extract_val(parts[4])
                                    
                                    exam = {
                                        "date": clean_date, 
                                        "time": str(time).strip(),
                                        "subject": subject_val, 
                                        "level": level_val,
                                        "levels_list": [clean_string_for_matching(l) for l in (level_val.split(' + ') if ' + ' in level_val else [level_val])],
                                        "professor": prof_val, 
                                        "halls": [{'name': h.strip(), 'type': _('غير محدد')} for h in halls_val.split('،') if h.strip()],
                                        "guards": [g.strip() for g in guards_val.split('،') if g.strip()]
                                    }
                                    final_schedule[exam['date']][exam['time']].append(exam)
        
        setting_sched = ExamSetting.query.filter_by(key='published_exam_schedule', tenant_id=tenant_id).first()
        value_str = json.dumps(final_schedule)
        if setting_sched: setting_sched.value = value_str
        else: db.session.add(ExamSetting(key='published_exam_schedule', value=value_str, tenant_id=tenant_id))

        setting_pub = ExamSetting.query.filter_by(key='is_exam_published', tenant_id=tenant_id).first()
        if setting_pub: setting_pub.value = '1'
        else: db.session.add(ExamSetting(key='is_exam_published', value='1', tenant_id=tenant_id))

        db.session.commit()
        return jsonify({"success": True, "message": _("تم استيراد الجدول النهائي ونشره للأساتذة بنجاح!"), "schedule": final_schedule})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500

