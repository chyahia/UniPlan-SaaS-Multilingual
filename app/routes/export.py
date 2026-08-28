# Copyright (c) 2026 Chaib Yahia. All rights reserved.
# This software is licensed under the CC BY-NC 4.0 License. Commercial use is strictly prohibited.
from flask import Blueprint, request, send_file, jsonify
import io
import re
from collections import defaultdict
import pandas as pd

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement # ضروري لاتجاه اليمين لليسار (RTL) في الجداول

from openpyxl.styles import Border, Side, Font, PatternFill, Alignment

from app.database import db, Teacher, Course
from flask import session

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from flask_babel import _ # ✨ استيراد دالة الترجمة لرسائل النظام
import traceback
from openpyxl.utils import get_column_letter

export_bp = Blueprint('export', __name__)

# ================== قاموس الترجمة الموحد (للملفات المصدرة) ==================
TRANSLATIONS = {
    'ar': {
        'time': 'الوقت', 'day': 'اليوم', 'teacher': 'الأستاذ', 'course': 'المادة', 'nature': 'طبيعة المادة', 'room': 'القاعة',
        'level': 'المستوى', 'empty': 'فارغ', 'no_teacher': 'بدون أستاذ', 'no_room': 'بدون قاعة',
        'days': {'السبت': 'السبت', 'الأحد': 'الأحد', 'الإثنين': 'الإثنين', 'الثلاثاء': 'الثلاثاء', 'الأربعاء': 'الأربعاء', 'الخميس': 'الخميس', 'الجمعة': 'الجمعة'},
        'nature_map': {'محاضرة': 'محاضرة', 'أعمال موجهة': 'أعمال موجهة', 'أعمال تطبيقية': 'أعمال تطبيقية'},
        'levels_map': {"Bachelor 1": "ليسانس 1", "Bachelor 2": "ليسانس 2", "Bachelor 3": "ليسانس 3", "Master 1": "ماستر 1", "Master 2": "ماستر 2"}
    },
    'en': {
        'time': 'Time', 'day': 'Day', 'teacher': 'Professor', 'course': 'Course', 'nature': 'Nature', 'room': 'Room',
        'level': 'Level', 'empty': 'Empty', 'no_teacher': 'No Professor', 'no_room': 'No Room',
        'days': {'السبت': 'Saturday', 'الأحد': 'Sunday', 'الإثنين': 'Monday', 'الثلاثاء': 'Tuesday', 'الأربعاء': 'Wednesday', 'الخميس': 'Thursday', 'الجمعة': 'Friday'},
        'nature_map': {'محاضرة': 'Lecture', 'أعمال موجهة': 'Tutorial (TD)', 'أعمال تطبيقية': 'Practical (TP)'},
        'levels_map': {"Bachelor 1": "Bachelor 1", "Bachelor 2": "Bachelor 2", "Bachelor 3": "Bachelor 3", "Master 1": "Master 1", "Master 2": "Master 2"}
    },
    'fr': {
        'time': 'Heure', 'day': 'Jour', 'teacher': 'Professeur', 'course': 'Module', 'nature': 'Nature', 'room': 'Salle',
        'level': 'Niveau', 'empty': 'Vide', 'no_teacher': 'Sans Professeur', 'no_room': 'Sans Salle',
        'days': {'السبت': 'Samedi', 'الأحد': 'Dimanche', 'الإثنين': 'Lundi', 'الثلاثاء': 'Mardi', 'الأربعاء': 'Mercredi', 'الخميس': 'Jeudi', 'الجمعة': 'Vendredi'},
        'nature_map': {'محاضرة': 'Cours', 'أعمال موجهة': 'TD', 'أعمال تطبيقية': 'TP'},
        'levels_map': {"Bachelor 1": "Licence 1", "Bachelor 2": "Licence 2", "Bachelor 3": "Licence 3", "Master 1": "Master 1", "Master 2": "Master 2"}
    }
}

# ================== دوال مساعدة لتنسيق الوورد والإكسل ==================
def create_word_document_with_table(doc, title, headers, data, lang='ar'):
    """دالة مساعدة لرسم جداول الوورد مع التحكم باتجاه اللغة"""
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if lang == 'ar':
        pPr = heading._element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
        for run in heading.runs:
            run.font.rtl = True
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    if lang == 'ar':
        tblPr = table._tbl.tblPr
        if tblPr is not None:
            bidiVisual = OxmlElement('w:bidiVisual')
            tblPr.append(bidiVisual)
        
    def set_cell_format(cell):
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
            if lang == 'ar':
                pPr = p._element.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.append(bidi)
                for run in p.runs:
                    run.font.rtl = True 

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
        set_cell_format(hdr_cells[i])
        
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            set_cell_format(row_cells[i])
            
    doc.add_page_break()

def process_and_format_sheet(writer, df, sheet_name, lang='ar'):
    """دالة مساعدة لتنسيق جداول الإكسل باحترافية مع دعم تعدد اللغات"""
    df.to_excel(writer, sheet_name=sheet_name)
    worksheet = writer.sheets[sheet_name]
    
    # تحديد اتجاه الشيت حسب اللغة
    is_rtl = (lang == 'ar')
    worksheet.sheet_view.rightToLeft = is_rtl 
    reading_order = 2 if is_rtl else 1
    
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
    center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True, readingOrder=reading_order) 

    # تنسيق الصف الأول (رؤوس الجدول)
    for cell in worksheet[1]:
        cell.font = header_font
        cell.fill = header_fill
        cell.border = thin_border
        cell.alignment = center_alignment
    worksheet.row_dimensions[1].height = 35

    # تنسيق باقي الخلايا
    for row in worksheet.iter_rows(min_row=2):
        # 🚀 الاحتواء التلقائي لارتفاع الصف ليتناسب مع الأسطر المتعددة
        worksheet.row_dimensions[row[0].row].height = None 
        for cell in row:
            cell.border = thin_border
            cell.alignment = center_alignment
            
    # ضبط عرض الأعمدة
    worksheet.column_dimensions['A'].width = 18 
    # 🚀 الحل الآمن والموثوق لتحويل أرقام الأعمدة إلى حروف
    for col_idx in range(2, len(df.columns) + 2):
        col_letter = get_column_letter(col_idx) 
        worksheet.column_dimensions[col_letter].width = 30

# =====================================================================
# 1. تصدير جداول المستويات (Word)
# =====================================================================
@export_bp.route('/api/export/word/all-levels', methods=['POST'])
def export_all_levels_word():
    data = request.get_json()
    schedules_by_level, days, slots = data.get('schedule'), data.get('days', []), data.get('slots', [])
    lang = data.get('lang', 'ar')
    
    if not all([schedules_by_level, days, slots]):
        # ✨ تغليف رسالة الخطأ بـ _
        return jsonify({"error": _("بيانات التصدير غير كاملة")}), 400

    t = TRANSLATIONS.get(lang, TRANSLATIONS['ar'])
    translated_days = [t['days'].get(d, d) for d in days]
    headers = [t['time']] + translated_days

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin_size = Cm(0.5)
    section.top_margin = margin_size; section.bottom_margin = margin_size
    section.left_margin = margin_size; section.right_margin = margin_size

    for level, grid_data in schedules_by_level.items():
        processed_data = []
        for i, slot_name in enumerate(slots):
            row_content = [slot_name]
            for j in range(len(days)):
                cell_text = "\n".join([f"{lec.get('name', '')}\n{lec.get('teacher_name', '')}\n{lec.get('room', '')}".strip() for lec in grid_data[j][i]])
                row_content.append(cell_text)
            processed_data.append(row_content)
        
        sheet_name = t['levels_map'].get(level, level)
        create_word_document_with_table(doc, sheet_name, headers, processed_data, lang)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    filename = 'Schedules_Levels.docx' if lang == 'en' else ('Emplois_Niveaux.docx' if lang == 'fr' else 'جداول_المستويات.docx')
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=filename)

# =====================================================================
# 2. تصدير جداول الأساتذة (Word)
# =====================================================================
@export_bp.route('/api/export/word/all-professors', methods=['POST'])
def export_all_professors_word():
    data = request.get_json()
    schedules_by_prof = data.get('prof_schedules') or data.get('schedule')
    days = data.get('days', [])
    slots = data.get('slots', [])
    lang = data.get('lang', 'ar')

    if not all([schedules_by_prof, days, slots]):
        # ✨ تغليف رسالة الخطأ
        return jsonify({"error": _("بيانات التصدير غير كاملة")}), 400

    t = TRANSLATIONS.get(lang, TRANSLATIONS['ar'])
    translated_days = [t['days'].get(d, d) for d in days]
    headers = [t['time']] + translated_days

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    margin_size = Cm(0.5)
    section.top_margin = margin_size; section.bottom_margin = margin_size
    section.left_margin = margin_size; section.right_margin = margin_size

    for prof_name, grid_data in sorted(schedules_by_prof.items()):
        processed_data = []
        for i, slot_name in enumerate(slots):
            row_content = [slot_name]
            for j in range(len(days)):
                lvl_lbl = t['level']
                cell_texts = [f"{lec.get('name', '')}\n{lvl_lbl}: {t['levels_map'].get(lec.get('level', ''), lec.get('level', ''))}\n{lec.get('room', '')}".strip() for lec in grid_data[j][i]]
                row_content.append("\n".join(cell_texts))
            processed_data.append(row_content)
        
        create_word_document_with_table(doc, prof_name, headers, processed_data, lang)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    filename = 'Schedules_Professors.docx' if lang == 'en' else ('Emplois_Professeurs.docx' if lang == 'fr' else 'جداول_الأساتذة.docx')
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=filename)

# =====================================================================
# 3. تصدير القاعات الفارغة (Excel)
# =====================================================================
@export_bp.route('/api/export/free-rooms', methods=['POST'])
def export_free_rooms():
    data = request.get_json()
    free_rooms_grid = data.get('free_rooms') or data.get('schedule')
    days = data.get('days', [])
    slots = data.get('slots', [])
    lang = data.get('lang', 'ar')
    
    if not all([free_rooms_grid, days, slots]): 
        # ✨ تغليف رسالة الخطأ
        return jsonify({"error": _("بيانات التصدير غير كاملة")}), 400
    
    t = TRANSLATIONS.get(lang, TRANSLATIONS['ar'])
    translated_days = [t['days'].get(d, d) for d in days]
    sheet_name = 'Free Rooms' if lang == 'en' else ('Salles Libres' if lang == 'fr' else 'القاعات الشاغرة')
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        processed_data = [["\n".join(free_rooms_grid[j][i]) for j in range(len(days))] for i in range(len(slots))]
        df = pd.DataFrame(processed_data, index=slots, columns=translated_days)
        process_and_format_sheet(writer, df, sheet_name, lang)
    
    output.seek(0)
    filename = 'Free_Rooms.xlsx' if lang == 'en' else ('Salles_Libres.xlsx' if lang == 'fr' else 'جدول_القاعات_الشاغرة.xlsx')
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', as_attachment=True, download_name=filename)

# =====================================================================
# 4. تصدير العبء البيداغوجي (Excel)
# =====================================================================
@export_bp.route('/api/export/teaching-load', methods=['GET'])
def export_teaching_load():
    try:
        tenant_id = session.get('tenant_id')
        lang = request.args.get('lang', 'ar')
        t = TRANSLATIONS.get(lang, TRANSLATIONS['ar'])
        
        teachers_list = Teacher.query.filter_by(tenant_id=tenant_id).all()
        teacher_map = {t_obj.id: t_obj.name for t_obj in teachers_list}
        courses_raw = Course.query.filter(Course.teacher_id != None, Course.tenant_id == tenant_id).all()
        
        courses_by_teacher = defaultdict(list)
        for c in courses_raw:
            teacher_name = teacher_map.get(c.teacher_id)
            if not teacher_name: continue
            course_dict = {
                'id': c.id, 'name': c.name, 'division': c.division, 'specialization': c.specialization,
                'levels': [t['levels_map'].get(l.name, l.name) for l in c.levels]
            }
            courses_by_teacher[teacher_name].append(course_dict)

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            workbook = writer.book
            if 'Sheet' in workbook.sheetnames:
                workbook.remove(workbook['Sheet'])
            sheet_title = 'Teaching Load' if lang == 'en' else ('Charge Pédagogique' if lang == 'fr' else 'العبء البيداغوجي')
            worksheet = workbook.create_sheet(sheet_title, 0)
            
            worksheet.sheet_view.rightToLeft = (lang == 'ar')
            reading_order = 2 if lang == 'ar' else 1

            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True, readingOrder=reading_order)
            cell_alignment = Alignment(horizontal='right' if lang == 'ar' else 'left', vertical='top', wrap_text=True, readingOrder=reading_order)
            merged_alignment = Alignment(horizontal='right' if lang == 'ar' else 'left', vertical='center', wrap_text=True, readingOrder=reading_order)
            banded_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")

            if lang == 'en':
                headers = ['No', 'Last Name', 'First Name', 'Rank', 'Degree', 'Specialty', 'Level', 'Type', 'Division', 'Specialization', 'Course Name', 'Department', 'Faculty', 'Hours']
            elif lang == 'fr':
                headers = ['N°', 'Nom', 'Prénom', 'Grade', 'Diplôme', 'Spécialité', 'Niveau', 'Type', 'Filière', 'Spécialisation', 'Module', 'Département', 'Faculté', 'Volume Horaire']
            else:
                headers = ['الرقم', 'اللقب', 'الاسم', 'الرتبة', 'الشهادة', 'تخصص الشهادة', 'المستوى', 'نوع المادة', 'الشعبة', 'التخصص', 'اسم المادة', 'القسم', 'الكلية', 'الحجم الساعي']

            for col_num, header_title in enumerate(headers, 1):
                cell = worksheet.cell(row=1, column=col_num)
                cell.value = header_title
                cell.font = header_font; cell.fill = header_fill; cell.border = thin_border; cell.alignment = header_alignment

            worksheet.column_dimensions['A'].width = 5
            worksheet.column_dimensions['B'].width = 30
            for col_letter in ['C', 'D', 'E', 'F']: worksheet.column_dimensions[col_letter].width = 15
            for col_letter in ['G', 'H']: worksheet.column_dimensions[col_letter].width = 15
            for col_letter in ['I', 'J']: worksheet.column_dimensions[col_letter].width = 25
            worksheet.column_dimensions['K'].width = 45
            for col_letter in ['L', 'M', 'N']: worksheet.column_dimensions[col_letter].width = 15

            current_row = 2
            professor_number = 1
            for teacher_name in sorted(courses_by_teacher.keys()):
                courses = courses_by_teacher[teacher_name]
                total_rows_for_teacher = sum(len(c.get('levels', [])) for c in courses)
                if total_rows_for_teacher == 0: continue
                
                teacher_start_row = current_row
                is_banded_row = (professor_number % 2 == 0)

                for course in courses:
                    for level_name in course.get('levels', []):
                        course_name_original = course.get('name', '')
                        course_type_ar = 'أعمال موجهة' 
                        if '[مح]' in course_name_original: course_type_ar = 'محاضرة'
                        elif '[أت]' in course_name_original: course_type_ar = 'أعمال تطبيقية'
                        
                        course_type_translated = t['nature_map'].get(course_type_ar, course_type_ar)
                        division = course.get('division') or ''
                        specialization = course.get('specialization') or ''

                        data_row = [level_name, course_type_translated, division, specialization, course_name_original]
                        for col_offset, value in enumerate(data_row):
                            worksheet.cell(row=current_row, column=7 + col_offset, value=value)
                        current_row += 1

                end_row = current_row - 1
                for r in range(teacher_start_row, end_row + 1):
                    worksheet.row_dimensions[r].height = None
                    for c in range(1, 15):
                        cell = worksheet.cell(row=r, column=c)
                        if c >= 7 and c <= 11: cell.alignment = cell_alignment
                        else: cell.alignment = merged_alignment
                        cell.border = thin_border
                        if is_banded_row: cell.fill = banded_fill

                if total_rows_for_teacher > 0:
                    cell_a = worksheet.cell(row=teacher_start_row, column=1); cell_a.value = professor_number; cell_a.font = Font(bold=True)
                    cell_b = worksheet.cell(row=teacher_start_row, column=2); cell_b.value = teacher_name; cell_b.font = Font(bold=True)
                    if total_rows_for_teacher > 1:
                        worksheet.merge_cells(start_row=teacher_start_row, start_column=1, end_row=end_row, end_column=1)
                        worksheet.merge_cells(start_row=teacher_start_row, start_column=2, end_row=end_row, end_column=2)
                        for col in [3, 4, 5, 6, 12, 13, 14]:
                            worksheet.merge_cells(start_row=teacher_start_row, start_column=col, end_row=end_row, end_column=col)
                professor_number += 1

        output.seek(0)
        filename = 'Teaching_Load.xlsx' if lang == 'en' else ('Charge_Pedagogique.xlsx' if lang == 'fr' else 'العبء_البيداغوجي_للأساتذة.xlsx')
        return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', as_attachment=True, download_name=filename)
    except Exception as e:
        import traceback
        traceback.print_exc()
        # ✨ تغليف رسالة الخطأ المتغيرة
        return jsonify({"error": _("فشل إنشاء الملف: {error}").format(error=str(e))}), 500
    
# =====================================================================
# 5. تصدير القائمة الشاملة (Excel - Flat Table)
# =====================================================================
@export_bp.route('/api/export/comprehensive-list', methods=['POST'])
def export_comprehensive_list():
    if 'tenant_id' not in session:
        # ✨ تغليف رسالة الخطأ
        return jsonify({"error": _("غير مصرح")}), 403

    data = request.get_json()
    schedule = data.get('schedule', {})
    days = data.get('days', [])
    slots = data.get('slots', [])
    lang = data.get('lang', 'ar')
    
    if not all([schedule, days, slots]): 
        # ✨ تغليف رسالة الخطأ
        return jsonify({"error": _("بيانات التصدير غير كاملة")}), 400
    
    t = TRANSLATIONS.get(lang, TRANSLATIONS['ar'])
    
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        if not schedule:
            pd.DataFrame([t['empty']]).to_excel(writer, sheet_name=t['empty'], index=False)
        
        for level, grid in schedule.items():
            safe_title = str(t['levels_map'].get(level, level)).replace("/", "-").replace("\\", "-").replace("*", "").replace("?", "")[:31]
            flat_data = []
            
            for d_idx, day_name in enumerate(days):
                translated_day = t['days'].get(day_name, day_name)
                for s_idx, slot_name in enumerate(slots):
                    lectures = grid[d_idx][s_idx] if d_idx < len(grid) and s_idx < len(grid[d_idx]) else []
                    
                    for lec in lectures:
                        course_name_full = lec.get('name', '')
                        teacher_name = lec.get('teacher_name') or t['no_teacher']
                        room_name = lec.get('room') or t['no_room']
                        
                        course_type_ar = 'أعمال موجهة' 
                        if '[مح]' in course_name_full: course_type_ar = 'محاضرة'
                        elif '[أت]' in course_name_full: course_type_ar = 'أعمال تطبيقية'
                        course_type = t['nature_map'].get(course_type_ar, course_type_ar)
                        
                        clean_course_name = course_name_full.replace('[مح]', '').replace('[أم]', '').replace('[أت]', '').strip()
                        if not clean_course_name: clean_course_name = course_name_full
                        
                        flat_data.append([translated_day, slot_name, teacher_name, clean_course_name, course_type, room_name])
            
            if not flat_data:
                flat_data = [["-", "-", "-", "-", "-", "-"]]
                
            headers = [t['day'], t['time'], t['teacher'], t['course'], t['nature'], t['room']]
            df = pd.DataFrame(flat_data, columns=headers)
            df.to_excel(writer, sheet_name=safe_title, index=False)
            
            worksheet = writer.sheets[safe_title]
            worksheet.sheet_view.rightToLeft = (lang == 'ar')
            reading_order = 2 if lang == 'ar' else 1
            
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
            rtl_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True, readingOrder=reading_order)
            
            for cell in worksheet[1]:
                cell.font = header_font; cell.fill = header_fill; cell.border = thin_border; cell.alignment = rtl_alignment
            
            for row in worksheet.iter_rows(min_row=2):
                for cell in row:
                    cell.border = thin_border; cell.alignment = rtl_alignment
            
            worksheet.column_dimensions['A'].width = 15 
            worksheet.column_dimensions['B'].width = 15 
            worksheet.column_dimensions['C'].width = 25 
            worksheet.column_dimensions['D'].width = 30 
            worksheet.column_dimensions['E'].width = 18 
            worksheet.column_dimensions['F'].width = 15 
            worksheet.auto_filter.ref = worksheet.dimensions
            
    output.seek(0)
    filename = 'Comprehensive_List.xlsx' if lang == 'en' else ('Liste_Globale.xlsx' if lang == 'fr' else 'القائمة_الشاملة_للجداول.xlsx')
    return send_file(output, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', as_attachment=True, download_name=filename)


# =====================================================================
# 6. تصدير جداول القاعات والمدرجات (Excel)
# =====================================================================
@export_bp.route('/api/export/excel/rooms-schedule', methods=['POST'])
def export_rooms_schedule_excel():
    # 🛡️ 1. التغليف الأمني: التأكد من الصلاحيات
    tenant_id = session.get('tenant_id')
    if not tenant_id:
        return jsonify({"error": "غير مصرح بالدخول أو انتهت الجلسة"}), 403

    # 🛡️ 2. تغليف الأخطاء: لمنع انهيار السيرفر (Try-Except)
    try:
        data = request.get_json() or {}
        schedule = data.get('schedule', {})
        days = data.get('days', [])
        slots = data.get('slots', [])
        lang = data.get('lang', 'ar') 
        
        if not all([schedule, days, slots]): 
            return jsonify({"error": "بيانات التصدير غير كاملة"}), 400

        # ✨ قاموس الترجمة للأيام والنصوص داخل الخلايا
        trans = {
            'ar': {
                'empty': "لا توجد قاعات مستخدمة في هذا الجدول", 'empty_sheet': "فارغ",
                'lec': "[محاضرة]", 'td': "[أعمال موجهة]", 'tp': "[أعمال تطبيقية]",
                'no_teacher': "بدون أستاذ",
                'days_map': {'الأحد': 'الأحد', 'الإثنين': 'الإثنين', 'الاثنين': 'الإثنين', 'الثلاثاء': 'الثلاثاء', 'الأربعاء': 'الأربعاء', 'الخميس': 'الخميس', 'الجمعة': 'الجمعة', 'السبت': 'السبت'}
            },
            'en': {
                'empty': "No rooms used in this schedule", 'empty_sheet': "Empty",
                'lec': "[Lecture]", 'td': "[TD]", 'tp': "[TP]",
                'no_teacher': "No Teacher",
                'days_map': {'الأحد': 'Sunday', 'الإثنين': 'Monday', 'الاثنين': 'Monday', 'الثلاثاء': 'Tuesday', 'الأربعاء': 'Wednesday', 'الخميس': 'Thursday', 'الجمعة': 'Friday', 'السبت': 'Saturday'}
            },
            'fr': {
                'empty': "Aucune salle utilisée dans cet emploi", 'empty_sheet': "Vide",
                'lec': "[Cours]", 'td': "[TD]", 'tp': "[TP]",
                'no_teacher': "Sans Professeur",
                'days_map': {'الأحد': 'Dimanche', 'الإثنين': 'Lundi', 'الاثنين': 'Lundi', 'الثلاثاء': 'Mardi', 'الأربعاء': 'Mercredi', 'الخميس': 'Jeudi', 'الجمعة': 'Vendredi', 'السبت': 'Samedi'}
            }
        }
        t = trans.get(lang, trans['ar'])

        translated_days = [t['days_map'].get(d, d) for d in days]
        
        # الهندسة العكسية: تجميع البيانات بناءً على القاعات
        room_schedules = defaultdict(lambda: [[ [] for _ in slots ] for _ in days])
        
        for level_name, lvl_grid in schedule.items():
            for d_idx, day_slots in enumerate(lvl_grid):
                for s_idx, lectures in enumerate(day_slots):
                    for lec in lectures:
                        room = lec.get('room')
                        if room and str(room).strip() not in ['بدون قاعة', 'None', '-', '']:
                            lec_copy = lec.copy()
                            lec_copy['level'] = level_name
                            room_schedules[str(room).strip()][d_idx][s_idx].append(lec_copy)

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            if not room_schedules:
                pd.DataFrame([t['empty']]).to_excel(writer, sheet_name=t['empty_sheet'], index=False)
            else:
                for room_name in sorted(room_schedules.keys()):
                    safe_title = str(room_name).replace("/", "-").replace("\\", "-").replace("*", "").replace("?", "")[:31]
                    grid = room_schedules[room_name]
                    
                    processed_data = []
                    for s_idx in range(len(slots)):
                        row_content = []
                        for d_idx in range(len(days)):
                            cell_lectures = grid[d_idx][s_idx]
                            if cell_lectures:
                                cell_texts = []
                                for l in cell_lectures:
                                    ctype = ""
                                    if '[مح]' in l.get('name', ''): ctype = t['lec']
                                    elif '[أت]' in l.get('name', ''): ctype = t['tp']
                                    else: ctype = t['td']
                                    
                                    clean_name = l.get('name', '').replace('[مح]', '').replace('[أم]', '').replace('[أت]', '').strip()
                                    
                                    teacher_val = l.get('teacher_name', '')
                                    if not teacher_val or teacher_val == 'بدون أستاذ':
                                        teacher_val = t['no_teacher']
                                        
                                    text = f"📚 {clean_name} {ctype}\n🎓 {l.get('level', '')}\n👨‍🏫 {teacher_val}"
                                    cell_texts.append(text)
                            
                                row_content.append("\n\n".join(cell_texts))
                            else:
                                row_content.append("-")
                        processed_data.append(row_content)
                    
                    df = pd.DataFrame(processed_data, index=slots, columns=translated_days)
                    process_and_format_sheet(writer, df, safe_title, lang=lang)
        
        output.seek(0)
        dl_name = 'Rooms_Schedule.xlsx' if lang != 'ar' else 'جداول_القاعات_الزمنية.xlsx'
        
        return send_file(
            output, 
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 
            as_attachment=True, 
            download_name=dl_name
        )

    # 🛡️ 3. التقاط الأخطاء وإرسالها للواجهة
    except Exception as e:
        traceback.print_exc() # لطباعة الخطأ في شاشة السيرفر للمطور
        return jsonify({"error": f"حدث خطأ أثناء تصدير الملف: {str(e)}"}), 500