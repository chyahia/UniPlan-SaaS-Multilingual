# Copyright (c) 2026 Chaib Yahia. All rights reserved.
# This software is licensed under the CC BY-NC 4.0 License. Commercial use is strictly prohibited.
import io
import docx
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from flask_babel import _

# ================== قاموس الترجمة للامتحانات الاستدراكية ==================
RESIT_TRANSLATIONS = {
    'ar': {
        'title_levels': _('جداول الامتحانات الاستدراكية - حسب المستويات'),
        'title_teachers': _('جداول الامتحانات الاستدراكية - للأساتذة'),
        'level': _('المستوى:'),
        'time_period': _('الفترة الزمنية'),
        'unknown_subject': _('مادة غير معروفة'),
        'unspecified': _('غير محدد'),
        'subject': _('المادة:'),
        'subject_teacher': _('أستاذ المادة:'),
        'guarding': _('الحراسة:\n'),
        'location': _('المكان:'),
        'no_guards': _('بدون حراس'),
        'teacher': _('الأستاذ(ة):'),
        'assigned_subjects': _('المواد المسندة:'),
        'no_subject': _('بدون مادة'),
        'days_map': {'الأحد': _('الأحد'), 'الإثنين': _('الإثنين'), 'الثلاثاء': _('الثلاثاء'), 'الأربعاء': _('الأربعاء'), 'الخميس': _('الخميس'), 'الجمعة': _('الجمعة'), 'السبت': _('السبت')}
    },
    'en': {
        'title_levels': 'Resit Exams Schedule - By Levels',
        'title_teachers': 'Resit Exams Schedule - For Teachers',
        'level': 'Level:',
        'time_period': 'Time Slot',
        'unknown_subject': 'Unknown Subject',
        'unspecified': 'Not Specified',
        'subject': 'Subject:',
        'subject_teacher': 'Course Professor:',
        'guarding': 'Guarding:\n',
        'location': 'Location:',
        'no_guards': 'No Guards',
        'teacher': 'Professor:',
        'assigned_subjects': 'Assigned Courses:',
        'no_subject': 'No Course',
        'days_map': {'الأحد': 'Sunday', 'الإثنين': 'Monday', 'الثلاثاء': 'Tuesday', 'الأربعاء': 'Wednesday', 'الخميس': 'Thursday', 'الجمعة': 'Friday', 'السبت': 'Saturday'}
    },
    'fr': {
        'title_levels': 'Emplois des Examens de Rattrapage - Par Niveaux',
        'title_teachers': 'Emplois des Examens de Rattrapage - Pour Professeurs',
        'level': 'Niveau:',
        'time_period': 'Créneau Horaire',
        'unknown_subject': 'Module Inconnu',
        'unspecified': 'Non Spécifié',
        'subject': 'Module:',
        'subject_teacher': 'Professeur du Module:',
        'guarding': 'Surveillance:\n',
        'location': 'Lieu:',
        'no_guards': 'Sans Surveillants',
        'teacher': 'Professeur:',
        'assigned_subjects': 'Modules Assignés:',
        'no_subject': 'Sans Module',
        'days_map': {'الأحد': 'Dimanche', 'الإثنين': 'Lundi', 'الثلاثاء': 'Mardi', 'الأربعاء': 'Mercredi', 'الخميس': 'Jeudi', 'الجمعة': 'Vendredi', 'السبت': 'Samedi'}
    }
}

def setup_landscape_doc():
    doc = docx.Document()
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)
    return doc

def make_table_rtl(table):
    """قلب اتجاه الجدول ليكون من اليمين لليسار (يُستدعى للعربية فقط)"""
    tblPr = table._element.tblPr
    for b in tblPr.findall(qn('w:bidiVisual')): 
        tblPr.remove(b)
    bidiVisual = OxmlElement('w:bidiVisual')
    tblPr.append(bidiVisual)

def set_cell_background(cell, color_hex):
    """تلوين خلفية الخلية"""
    tcPr = cell._element.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
        
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    tcPr.append(shading_elm)

def format_paragraph(p, font_size=14, bold=False, align_center=False, lang='ar'):
    """دالة التنسيق الآمنة مع دعم الاتجاهات (RTL/LTR) بناءً على اللغة"""
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if lang == 'ar' else WD_ALIGN_PARAGRAPH.LEFT
    
    for run in p.runs:
        run.font.size = Pt(font_size)
        run.font.name = 'Arial'
        if bold:
            run.font.bold = True
            
        if lang == 'ar':
            rPr = run._element.get_or_add_rPr()
            
            rtl = OxmlElement('w:rtl')
            rtl.set(qn('w:val'), '1')
            rPr.append(rtl)
            
            cs = OxmlElement('w:cs')
            cs.set(qn('w:val'), '1')
            rPr.append(cs)
            
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(font_size * 2))
            rPr.append(szCs)
            
            if bold:
                bCs = OxmlElement('w:bCs')
                bCs.set(qn('w:val'), '1')
                rPr.append(bCs)

def generate_levels_word(db, distribution, lang='ar'):
    doc = setup_landscape_doc()
    t = RESIT_TRANSLATIONS.get(lang, RESIT_TRANSLATIONS['ar'])
    
    h = doc.add_heading(t['title_levels'], 0)
    format_paragraph(h, font_size=20, bold=True, align_center=True, lang=lang)
    
    levels = db.get('levels', [])
    schedule = db.get('schedule', {})
    
    days = list(schedule.keys())
    times_list = []
    for d in days:
        for slot in schedule[d]:
            if slot['time'] not in times_list:
                times_list.append(slot['time'])
                
    for level in levels:
        h_lvl = doc.add_heading(f"{t['level']} {level}", level=1)
        format_paragraph(h_lvl, font_size=16, bold=True, lang=lang)
        
        table = doc.add_table(rows=len(times_list) + 1, cols=len(days) + 1)
        table.style = 'Table Grid'
        if lang == 'ar':
            make_table_rtl(table)
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = t['time_period']
        set_cell_background(hdr_cells[0], "D9E2F3") 
        for p in hdr_cells[0].paragraphs:
            format_paragraph(p, font_size=15, bold=True, align_center=True, lang=lang)
            
        for i, day in enumerate(days):
            # ترجمة اسم اليوم فقط مع الإبقاء على التاريخ
            translated_day = day
            for ar_day, trans_day in t.get('days_map', {}).items():
                if ar_day in day:
                    translated_day = day.replace(ar_day, trans_day)
                    break
                    
            hdr_cells[i+1].text = translated_day
            set_cell_background(hdr_cells[i+1], "D9E2F3") 
            for p in hdr_cells[i+1].paragraphs:
                format_paragraph(p, font_size=15, bold=True, align_center=True, lang=lang)
            
        for r_idx, time_val in enumerate(times_list):
            row_cells = table.rows[r_idx + 1].cells
            
            row_cells[0].text = time_val
            set_cell_background(row_cells[0], "D9E2F3") 
            for p in row_cells[0].paragraphs:
                format_paragraph(p, font_size=14, bold=True, align_center=True, lang=lang)
            
            for c_idx, day in enumerate(days):
                cell = row_cells[c_idx + 1]
                day_data = distribution.get(day, {})
                if time_val in day_data:
                    levels_dict = day_data[time_val]
                    if level in levels_dict:
                        subject_name = levels_dict[level].get("subject", t['unknown_subject'])
                        sub_teachers = levels_dict[level].get("subject_teachers", [])
                        teacher_str = "، ".join(sub_teachers) if sub_teachers else t['unspecified']
                        
                        cell_text = f"{t['subject']} {subject_name}\n{t['subject_teacher']} {teacher_str}\n\n{t['guarding']}"
                        
                        for room, teachers in levels_dict[level]["rooms"].items():
                            cell_text += f"{t['location']} {room}\n"
                            cell_text += "، ".join(teachers) if teachers else t['no_guards']
                            cell_text += "\n"
                            
                        cell.text = cell_text.strip()
                
                for p in cell.paragraphs:
                    format_paragraph(p, font_size=14, bold=False, align_center=False, lang=lang)
        
        doc.add_page_break()
        
    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)
    return mem

def generate_teachers_word(db, distribution, lang='ar'):
    doc = setup_landscape_doc()
    t = RESIT_TRANSLATIONS.get(lang, RESIT_TRANSLATIONS['ar'])
    
    h = doc.add_heading(t['title_teachers'], 0)
    format_paragraph(h, font_size=20, bold=True, align_center=True, lang=lang)
    
    teachers = db.get('teachers', [])
    schedule = db.get('schedule', {})
    teacher_subjects = db.get('teacher_subjects', {})
    
    days = list(schedule.keys())
    times_list = []
    for d in days:
        for slot in schedule[d]:
            if slot['time'] not in times_list:
                times_list.append(slot['time'])
                
    for teacher in teachers:
        h_t = doc.add_heading(f"{t['teacher']} {teacher}", level=1)
        format_paragraph(h_t, font_size=16, bold=True, lang=lang)
        
        subs = teacher_subjects.get(teacher, [])
        if subs:
            p = doc.add_paragraph(f"{t['assigned_subjects']} {'، '.join(subs)}")
            format_paragraph(p, font_size=14, bold=False, lang=lang)
        
        table = doc.add_table(rows=len(times_list) + 1, cols=len(days) + 1)
        table.style = 'Table Grid'
        if lang == 'ar':
            make_table_rtl(table)
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = t['time_period']
        set_cell_background(hdr_cells[0], "D9E2F3")
        for p in hdr_cells[0].paragraphs:
            format_paragraph(p, font_size=15, bold=True, align_center=True, lang=lang)
            
        for i, day in enumerate(days):
            # ترجمة اسم اليوم فقط مع الإبقاء على التاريخ
            translated_day = day
            for ar_day, trans_day in t.get('days_map', {}).items():
                if ar_day in day:
                    translated_day = day.replace(ar_day, trans_day)
                    break
                    
            hdr_cells[i+1].text = translated_day
            set_cell_background(hdr_cells[i+1], "D9E2F3") 
            for p in hdr_cells[i+1].paragraphs:
                format_paragraph(p, font_size=15, bold=True, align_center=True, lang=lang)
            
        for r_idx, time_val in enumerate(times_list):
            row_cells = table.rows[r_idx + 1].cells
            
            row_cells[0].text = time_val
            set_cell_background(row_cells[0], "D9E2F3")
            for p in row_cells[0].paragraphs:
                format_paragraph(p, font_size=14, bold=True, align_center=True, lang=lang)
            
            for c_idx, day in enumerate(days):
                cell = row_cells[c_idx + 1]
                cell_text = ""
                
                day_data = distribution.get(day, {})
                if time_val in day_data:
                    for level, level_dict in day_data[time_val].items():
                        rooms_dict = level_dict.get("rooms", {})
                        subject_name = level_dict.get("subject", t['no_subject'])
                        
                        for room, assigned_teachers in rooms_dict.items():
                            if teacher in assigned_teachers:
                                cell_text += f"{t['subject']} {subject_name}\n{t['level']} {level}\n{t['location']} {room}\n\n"
                                
                cell.text = cell_text.strip()
                for p in cell.paragraphs:
                    format_paragraph(p, font_size=14, bold=False, align_center=False, lang=lang)
                    
        doc.add_page_break()
        
    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)
    return mem